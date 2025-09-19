import docx
import os
from docx.shared import Inches
# Import the new, correct, document-based AI function
from ai_analyzer import analyze_pdf_direct_non_stream

def create_docx_from_ai_data(structured_data, docx_path):
    """Creates a DOCX file from the structured JSON data returned by the AI."""
    doc = docx.Document()
    if not isinstance(structured_data, list):
        print(f"Error: AI returned unexpected data format: {structured_data}")
        doc.add_paragraph(f"AI分析失败，返回数据格式错误或API调用失败: {structured_data}")
        doc.save(docx_path)
        return

    for i, group in enumerate(structured_data):
        doc.add_heading(f'资料分析题 {i + 1}', level=1)
        doc.add_paragraph(group.get('material', ''))
        
        for question in group.get('questions', []):
            doc.add_paragraph(f"\n题目 {question.get('question_number', '')}:")
            doc.add_paragraph(question.get('question_text', '').strip())
            doc.add_paragraph(f"正确答案： {question.get('correct_answer', '')}")
            doc.add_paragraph(f"你的答案： {question.get('user_answer', '')}")
        
        doc.add_page_break()
    
    doc.save(docx_path)
    print(f"Successfully created DOCX file at '{docx_path}'")

def convert_pdf_with_ai_direct(pdf_file, output_docx):
    """Main function to drive the direct PDF AI conversion."""
    print("Step 1: Calling AI to analyze PDF directly...")
    structured_json = analyze_pdf_direct_non_stream(pdf_file)
    
    if structured_json and 'error' not in structured_json:
        print("Step 2: Creating DOCX from structured data...")
        create_docx_from_ai_data(structured_json, output_docx)
    else:
        print("Failed to get structured data from AI.")
        print(f"AI Response: {structured_json}")

if __name__ == '__main__':
    pdf_input = 'uploads/专项智能练习（资料分析） 2.pdf'
    docx_output = 'uploads/专项智能练习（资料分析）_ai_final_direct.docx'
    
    convert_pdf_with_ai_direct(pdf_input, docx_output)
