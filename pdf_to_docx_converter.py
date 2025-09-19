import fitz  # PyMuPDF
import docx
import os
import re
from docx.shared import Inches

def write_group_to_doc_offline(doc, group_content, group_number):
    """Writes a collected group to the DOCX, merging material text correctly."""
    if not group_content:
        return
    
    doc.add_heading(f'资料分析题 {group_number}', level=1)
    
    question_start_regex = re.compile(r'^\s*(\d{1,2})\.')
    material_parts = []
    questions_and_images = []
    
    in_material_section = True
    for content_type, content_data in group_content:
        if content_type == 'text' and question_start_regex.match(content_data.strip()):
            in_material_section = False
        
        if in_material_section:
            if content_type == 'text':
                material_parts.append(content_data.replace('\n', ' '))
        else:
            questions_and_images.append((content_type, content_data))

    if material_parts:
        full_material = " ".join(material_parts)
        doc.add_paragraph(full_material)

    for content_type, content_data in questions_and_images:
        if content_type == 'text':
            if content_data.strip().startswith("正确答案：") or content_data.strip().startswith("你的答案："):
                 p = doc.add_paragraph()
                 p.add_run(content_data).italic = True
            else:
                doc.add_paragraph(content_data)
        elif content_type == 'image':
            try:
                doc.add_picture(content_data, width=Inches(5.5))
            except Exception as e:
                print(f"Could not add image {content_data}: {e}")
    doc.add_page_break()

def offline_final_converter(pdf_path, docx_path, image_dir='temp_images_offline'):
    """The final offline version with the most robust logic."""
    os.makedirs(image_dir, exist_ok=True)
    pdf_document = fitz.open(pdf_path)
    doc = docx.Document()
    
    question_start_regex = re.compile(r'^\s*(\d{1,2})\.')
    answer_regex = re.compile(r'正确答案：|你的答案：')
    unwanted_text_patterns = ["本试卷由粉笔用户", "专项智能练习（资料分析）"]
    
    all_elements = []
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        page_width = page.rect.width
        
        # Use a detailed text extraction method
        blocks = page.get_text("dict", flags=fitz.TEXTFLAGS_TEXT)["blocks"]
        for block in blocks:
            if block['type'] == 0: # Text block
                for line in block['lines']:
                    line_text = "".join([span['text'] for span in line['spans']]).strip()
                    if line_text and not any(unwanted in line_text for unwanted in unwanted_text_patterns):
                        all_elements.append({'type': 'text', 'content': line_text})
        
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            bbox = page.get_image_bbox(img)
            if bbox and (bbox.width / page_width < 0.8):
                image_filename = f"{image_dir}/page{page_num+1}_img{img_index}.png"
                img_data = pdf_document.extract_image(xref)
                with open(image_filename, "wb") as img_file:
                    img_file.write(img_data["image"])
                all_elements.append({'type': 'image', 'content': image_filename})

    current_group = []
    group_counter = 1
    seen_question_in_group = False

    for el in all_elements:
        is_question = el['type'] == 'text' and question_start_regex.match(el['content'])
        is_answer = el['type'] == 'text' and answer_regex.search(el['content'])
        is_new_material = el['type'] == 'text' and not is_question and not is_answer

        if is_new_material and seen_question_in_group:
            write_group_to_doc_offline(doc, current_group, group_counter)
            group_counter += 1
            current_group = [ (el['type'], el['content']) ]
            seen_question_in_group = False
        else:
            if is_question:
                seen_question_in_group = True
            current_group.append((el['type'], el['content']))

    write_group_to_doc_offline(doc, current_group, group_counter)

    doc.save(docx_path)
    pdf_document.close()
    print(f"Successfully converted '{pdf_path}' using the final offline logic to '{docx_path}'")

if __name__ == '__main__':
    pdf_input_path = 'uploads/专项智能练习（资料分析） 2.pdf'
    docx_output_path = 'uploads/专项智能练习（资料分析）_offline_final.docx'
    
    offline_final_converter(pdf_input_path, docx_output_path)
